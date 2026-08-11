"""safe_docx.py — Word document reader + L0 sanitizer.

Extracts visible text from .docx files using python-docx, filtering out
hidden runs (w:vanish property) and near-white font color text, then runs
L0 sanitization before returning content.

Attack vectors specific to .docx:
  - w:vanish / run.font.hidden = True — text marked hidden in Word, invisible
    in the document but extracted by all parsers
  - Near-white font color — same idea as white text in PDFs
  - Revision history (deleted text) — tracked changes still in XML

Out of scope (v1):
  - Revision/tracked changes (deleted text in markup)
  - Comments and annotations
  - Embedded objects and macros
  - Theme color detection (theme colors could theoretically be white)
  - Password-protected documents

Usage (from Claude via Bash):
    python3 /path/to/safe_docx.py '/path/to/file.docx'

Exits non-zero on error. Output is clean plaintext on stdout.
"""
import sys
import os

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _THIS_DIR)

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print(
        "ERROR: python-docx not installed.\n"
        "Install with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)

from sanitize import sanitize, NO_CAP

try:
    from safe_input import scan_for_injection, provenance_route, resolve_desk
except ImportError:
    scan_for_injection = None; provenance_route = None; resolve_desk = None

_MAX_FILE_BYTES = 10_000_000  # 10 MB cap
_WHITE_THRESHOLD = 229        # RGB value above this (out of 255) = near-white


def _is_hidden_run(run) -> bool:
    """Return True if this run should be excluded from output."""
    # w:vanish — text hidden in Word (invisible in document view)
    try:
        if run.font.hidden is True:
            return True
    except Exception:
        pass

    # Near-white font color — RGBColor is bytes-like: rgb[0]=R, rgb[1]=G, rgb[2]=B
    try:
        rgb = run.font.color.rgb
        if rgb is not None:
            r, g, b = rgb[0], rgb[1], rgb[2]
            if r > _WHITE_THRESHOLD and g > _WHITE_THRESHOLD and b > _WHITE_THRESHOLD:
                return True
    except Exception:
        pass

    return False


def _extract_metadata(doc) -> str:
    """Extract and sanitize content-bearing core properties."""
    try:
        props = doc.core_properties
    except Exception:
        return ""

    fields = {
        "Title": getattr(props, "title", None),
        "Author": getattr(props, "author", None),
        "Subject": getattr(props, "subject", None),
        "Keywords": getattr(props, "keywords", None),
        "Description": getattr(props, "description", None),
    }

    parts = []
    for field, value in fields.items():
        if value and str(value).strip():
            clean = sanitize(str(value), max_len=200)
            if clean:
                parts.append(f"[{field}: {clean}]")
    return "\n".join(parts)


def extract_and_sanitize(path: str, desk: str = "root") -> str:
    """Extract visible text from a .docx file, filter hidden runs, run L0 sanitization.

    Returns clean plaintext. Raises RuntimeError on unrecoverable errors.
    """
    try:
        file_size = os.path.getsize(path)
    except OSError as e:
        raise RuntimeError(f"Cannot access file {path}: {e}") from e

    if file_size > _MAX_FILE_BYTES:
        raise RuntimeError(
            f"File too large: {file_size:,} bytes (max {_MAX_FILE_BYTES:,})."
        )

    try:
        doc = Document(path)
    except Exception as e:
        msg = str(e)
        if "not a zip" in msg.lower() or "badzip" in msg.lower() or "corrupt" in msg.lower():
            raise RuntimeError(f"Corrupt or invalid .docx file: {e}") from e
        raise RuntimeError(f"Failed to open .docx: {e}") from e

    # Metadata
    meta_text = _extract_metadata(doc)

    # Body — collect visible runs from all paragraphs
    body_parts = []
    for para in doc.paragraphs:
        para_chars = []
        for run in para.runs:
            if not _is_hidden_run(run):
                para_chars.append(run.text)
        para_text = "".join(para_chars).strip()
        if para_text:
            body_parts.append(para_text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    cell_chars = []
                    for run in para.runs:
                        if not _is_hidden_run(run):
                            cell_chars.append(run.text)
                    cell_text = "".join(cell_chars).strip()
                    if cell_text:
                        body_parts.append(cell_text)

    body_raw = "\n".join(body_parts)
    body_clean = sanitize(body_raw, max_len=NO_CAP)

    sections = []
    if meta_text:
        sections.append(meta_text)
    if body_clean:
        sections.append(body_clean)

    result = "\n\n".join(sections)

    # Heuristic injection scan — flags to stderr
    if scan_for_injection is not None:
        findings = scan_for_injection(result)
        if provenance_route is not None:
            provenance_route(desk, "file", result, item=path)   # on-path gate (W5): provenance + coverage breadcrumb + Sentinel verdict; format preserved in item=path
        if findings:
            print(f"[safe_docx] FLAGGED — {len(findings)} injection pattern(s) detected:", file=sys.stderr)
            for match, label in findings:
                print(f"  [{label}] \"{match}\"", file=sys.stderr)

    return result


if __name__ == "__main__":
    desk, rest = resolve_desk() if resolve_desk else ("root", sys.argv[1:])
    if not rest:
        print("Usage: python3 safe_docx.py [--desk <id>] '/path/to/file.docx'", file=sys.stderr)
        sys.exit(1)

    docx_path = rest[0]
    try:
        result = extract_and_sanitize(docx_path, desk=desk)
        print(result)
    except RuntimeError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
